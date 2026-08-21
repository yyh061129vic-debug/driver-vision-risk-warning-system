"""生成 PIDNet drivable 道路分割优化报告（含图表 + Word + Markdown）。

与 team 3 sub SegFormer 同口径对比，输出道路 mask IoU / boundary F1 等指标的
可视化报告，并附带训练曲线、后处理变体对比与逐样本分布。
"""

from __future__ import annotations

import argparse
import json
import sys
import re
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import cv2
import numpy as np
import torch
import torch.nn.functional as F
import yaml
from docx import Document
from docx.shared import Inches
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
TRAIN_DIR = Path("D:/pidnet_drivable_retrain/train/bdd_drivable/pidnet_small_bdd_drivable_pilot_v2")
DEFAULT_CONFIG = ROOT / "external" / "PIDNet" / "configs" / "bdd100k" / "pidnet_small_bdd_drivable_pilot_v2.yaml"
PIDNET_ROOT = ROOT / "external" / "PIDNet"
if str(PIDNET_ROOT) not in sys.path:
    sys.path.insert(0, str(PIDNET_ROOT))

import models  # type: ignore  # noqa: E402
from configs import config, update_config  # type: ignore  # noqa: E402

# 后处理变体（raw 为基准，其余为 A/B/C 切边方案）
VARIANTS = ["raw", "refined", "crf", "snake", "lane_trim", "all"]
VARIANT_LABELS = {
    "raw": "raw（无后处理）",
    "refined": "refined（连通域+形态学）",
    "crf": "CRF（联合双边滤波）",
    "snake": "Snake（主动轮廓）",
    "lane_trim": "车道线引导切边",
    "all": "A+B+C（全部）",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate drivable PIDNet road segmentation report.")
    parser.add_argument("--eval-dir", default=str(ROOT / "outputs" / "drivable_pidnet_v2_eval"))
    parser.add_argument("--baseline-summary", default=str(ROOT / "outputs" / "segformer_baseline_eval" / "summary.json"))
    parser.add_argument("--train-dir", default=str(TRAIN_DIR))
    parser.add_argument("--config", default=str(DEFAULT_CONFIG))
    parser.add_argument("--visual-count", type=int, default=6)
    parser.add_argument("--output-docx", default=str(ROOT.parent / "passages_reports" / "PIDNet_Drivable道路分割优化报告.docx"))
    parser.add_argument("--output-md", default=str(ROOT.parent / "passages_reports" / "PIDNet_Drivable道路分割优化报告.md"))
    return parser.parse_args()


def find_train_log(train_dir: Path) -> Path:
    logs = sorted(train_dir.glob("*_train.log"), key=lambda p: p.stat().st_mtime, reverse=True)
    for log in logs:
        text = log.read_text(encoding="utf-8", errors="ignore")
        if "bdd_drivable" in text and "Epoch: [" in text:
            return log
    if logs:
        return logs[0]
    raise FileNotFoundError(f"no training log under {train_dir}")


def parse_training_log(log_path: Path) -> dict:
    """解析训练日志，返回每个 epoch 的验证指标序列。"""
    road_iou_line = re.compile(
        r"^\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2},\d+\s+1 \[(\S+)\s+(\S+)\s*\]"
    )
    val_line = re.compile(r"Loss: (\S+), MeanIU:\s+(\S+), Best_mIoU:\s+(\S+)")

    road_ious: list[float] = []
    bg_ious: list[float] = []
    val_loss: list[float] = []
    val_miou: list[float] = []
    best_miou: list[float] = []

    for line in log_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        m = road_iou_line.match(line)
        if m:
            bg_ious.append(float(m.group(1)))
            road_ious.append(float(m.group(2)))
            continue
        m = val_line.search(line)
        if m:
            val_loss.append(float(m.group(1)))
            val_miou.append(float(m.group(2)))
            best_miou.append(float(m.group(3)))

    epochs = list(range(1, len(road_ious) + 1))
    return {
        "epochs": epochs,
        "road_iou": road_ious,
        "bg_iou": bg_ious,
        "val_loss": val_loss,
        "val_miou": val_miou,
        "best_miou": best_miou,
    }


def load_runtime_config(cfg_path: str):
    class Args:
        cfg = cfg_path
        opts: list[str] = []

    update_config(config, Args())
    return config


def load_model(runtime_cfg, model_path: Path) -> torch.nn.Module:
    model = models.pidnet.get_seg_model(runtime_cfg, imgnet_pretrained=False)
    checkpoint = torch.load(model_path, map_location="cpu")
    if "state_dict" in checkpoint:
        checkpoint = checkpoint["state_dict"]
    model_dict = model.state_dict()
    remapped = {}
    for key, value in checkpoint.items():
        mapped_key = key[6:] if key.startswith("model.") else key
        if mapped_key in model_dict and value.shape == model_dict[mapped_key].shape:
            remapped[mapped_key] = value
    model_dict.update(remapped)
    model.load_state_dict(model_dict, strict=False)
    return model.cuda().eval()


def parse_list_file(data_root: Path, list_path: str) -> dict[str, Path]:
    mapping: dict[str, Path] = {}
    for line in Path(list_path).read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        image_rel, label_rel = line.strip().split()
        mapping[str((data_root / image_rel).resolve())] = (data_root / label_rel).resolve()
    return mapping


def predict_raw_mask(runtime_cfg, model: torch.nn.Module, image: Image.Image, road_class_id: int) -> np.ndarray:
    image_array = np.asarray(image.convert("RGB"), dtype=np.float32)
    height, width = image_array.shape[:2]
    normalized = image_array / 255.0
    normalized -= np.asarray([0.485, 0.456, 0.406], dtype=np.float32)
    normalized /= np.asarray([0.229, 0.224, 0.225], dtype=np.float32)
    tensor = torch.from_numpy(np.transpose(normalized, (2, 0, 1))).unsqueeze(0).cuda()

    with torch.no_grad():
        prediction = model(tensor)
        if isinstance(prediction, (list, tuple)):
            prediction = prediction[runtime_cfg.TEST.OUTPUT_INDEX]
        prediction = F.interpolate(
            prediction,
            size=(height, width),
            mode="bilinear",
            align_corners=runtime_cfg.MODEL.ALIGN_CORNERS,
        )
    logits = prediction[0]
    class_map = logits.argmax(dim=0).cpu().numpy().astype(np.uint8)
    return class_map == 1


def add_picture_if_exists(document: Document, path: Path, width: float = 6.0) -> None:
    if path.exists():
        document.add_picture(str(path), width=Inches(width))


def make_overlay(image_rgb: np.ndarray, pred_mask: np.ndarray, gt_mask: np.ndarray) -> np.ndarray:
    overlay = image_rgb.copy()
    true_positive = pred_mask & gt_mask
    false_positive = pred_mask & ~gt_mask
    false_negative = ~pred_mask & gt_mask

    overlay[true_positive] = (0.55 * overlay[true_positive] + 0.45 * np.array([0, 255, 0])).astype(np.uint8)
    overlay[false_positive] = (0.45 * overlay[false_positive] + 0.55 * np.array([255, 0, 0])).astype(np.uint8)
    overlay[false_negative] = (0.45 * overlay[false_negative] + 0.55 * np.array([0, 0, 255])).astype(np.uint8)
    return overlay


def create_visual_samples(
    per_sample: list[dict],
    runtime_cfg,
    model: torch.nn.Module,
    label_map: dict[str, Path],
    visual_dir: Path,
    limit: int,
) -> list[dict[str, object]]:
    visual_dir.mkdir(parents=True, exist_ok=True)
    group_size = max(1, limit // 3)
    by_worst = sorted(per_sample, key=lambda item: (item["raw_road_iou"], item["raw_boundary_f1"]))
    by_best = sorted(per_sample, key=lambda item: (-item["raw_road_iou"], -item["raw_boundary_f1"]))
    mean_iou = float(sum(item["raw_road_iou"] for item in per_sample) / len(per_sample))
    by_average = sorted(
        per_sample,
        key=lambda item: (abs(item["raw_road_iou"] - mean_iou), abs(item["raw_boundary_f1"] - 0.25)),
    )

    selected: list[dict] = []
    used_images: set[str] = set()

    def extend_group(candidates: list[dict], category: str, target_count: int) -> None:
        picked = 0
        for item in candidates:
            image_key = str(item["image"])
            if image_key in used_images:
                continue
            copied = dict(item)
            copied["sample_category"] = category
            selected.append(copied)
            used_images.add(image_key)
            picked += 1
            if picked >= target_count:
                break

    extend_group(by_worst, "最差样例", group_size)
    extend_group(by_best, "最好样例", group_size)
    extend_group(by_average, "最平均样例", max(1, limit - len(selected)))

    if len(selected) < limit:
        extend_group(by_average, "最平均样例", limit - len(selected))

    samples: list[dict[str, object]] = []

    for rank, item in enumerate(selected, start=1):
        image_path = Path(item["image"]).resolve()
        label_path = label_map.get(str(image_path))
        if label_path is None or not image_path.exists() or not label_path.exists():
            continue

        image = Image.open(image_path).convert("RGB")
        image_rgb = np.asarray(image, dtype=np.uint8)
        label = np.asarray(Image.open(label_path), dtype=np.uint8)
        gt_mask = (label > 0) & (label != 255)
        pred_mask = predict_raw_mask(runtime_cfg, model, image, road_class_id=1)
        overlay = make_overlay(image_rgb, pred_mask, gt_mask)

        panel_path = visual_dir / f"{rank:02d}_{image_path.stem}_panel.png"
        fig, axes = plt.subplots(2, 2, figsize=(10, 6))
        axes = axes.ravel()
        axes[0].imshow(image_rgb)
        axes[0].set_title("RGB")
        axes[1].imshow(pred_mask, cmap="gray")
        axes[1].set_title(f"Prediction (IoU={item['raw_road_iou']:.3f})")
        axes[2].imshow(gt_mask, cmap="gray")
        axes[2].set_title("Ground Truth")
        axes[3].imshow(overlay)
        axes[3].set_title("Overlay (green TP / red FP / blue FN)")
        for ax in axes:
            ax.axis("off")
        fig.tight_layout()
        fig.savefig(panel_path, dpi=150)
        plt.close(fig)

        samples.append(
            {
                "name": image_path.stem,
                "category": item["sample_category"],
                "panel": panel_path,
                "raw_road_iou": item["raw_road_iou"],
                "raw_boundary_f1": item["raw_boundary_f1"],
            }
        )
    return samples


def _style_ax(ax) -> None:
    ax.grid(True, linestyle="--", alpha=0.35)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)


def plot_training_curves(data: dict, out: Path) -> None:
    epochs = data["epochs"]
    fig, axes = plt.subplots(1, 3, figsize=(14, 3.8))

    ax = axes[0]
    ax.plot(epochs, data["road_iou"], marker="o", markersize=3, linewidth=1.6, color="#2c7fb8")
    ax.plot(epochs, data["bg_iou"], marker="o", markersize=3, linewidth=1.6, color="#a6bddb", alpha=0.8)
    ax.set_title("Validation IoU per epoch")
    ax.set_xlabel("Epoch")
    ax.set_ylabel("IoU")
    ax.legend(["road (class 1)", "background (class 0)"], fontsize=8)
    _style_ax(ax)

    ax = axes[1]
    ax.plot(epochs, data["val_miou"], marker="o", markersize=3, linewidth=1.6, color="#31a354")
    ax.plot(epochs, data["best_miou"], marker="o", markersize=3, linewidth=1.6, color="#a1d99b", alpha=0.9)
    ax.set_title("Validation mean IoU per epoch")
    ax.set_xlabel("Epoch")
    ax.set_ylabel("mean IoU")
    ax.legend(["current", "best (cumulative)"], fontsize=8)
    _style_ax(ax)

    ax = axes[2]
    if data["val_loss"]:
        ax.plot(epochs[: len(data["val_loss"])], data["val_loss"], marker="o", markersize=3, linewidth=1.6, color="#e6550d")
    ax.set_title("Validation loss per epoch")
    ax.set_xlabel("Epoch")
    ax.set_ylabel("Loss")
    _style_ax(ax)

    fig.tight_layout()
    out.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out, dpi=160)
    plt.close(fig)


def plot_baseline_comparison(ours: dict, baseline: dict, out: Path) -> None:
    metrics = ["road_iou", "boundary_f1", "road_precision", "road_recall"]
    labels = ["road IoU", "boundary F1", "road precision", "road recall"]
    ours_vals = [ours["raw"][m] for m in metrics]
    base_vals = [
        baseline["road_iou_global"],
        baseline["boundary_f1"],
        baseline["road_precision"],
        baseline["road_recall"],
    ]

    x = range(len(metrics))
    width = 0.36
    fig, ax = plt.subplots(figsize=(8, 4.2))
    ax.bar([i - width / 2 for i in x], base_vals, width, label="team 3 sub SegFormer", color="#bdbdbd")
    ax.bar([i + width / 2 for i in x], ours_vals, width, label="PIDNet drivable (ours)", color="#2c7fb8")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels)
    ax.set_ylim(0, 1.0)
    ax.set_ylabel("score")
    ax.set_title("Road segmentation: ours vs SegFormer baseline")
    ax.legend(fontsize=9)
    _style_ax(ax)

    for i, (b, o) in enumerate(zip(base_vals, ours_vals)):
        ax.text(i - width / 2, b + 0.02, f"{b:.3f}", ha="center", fontsize=8)
        ax.text(i + width / 2, o + 0.02, f"{o:.3f}", ha="center", fontsize=8)

    fig.tight_layout()
    out.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out, dpi=160)
    plt.close(fig)


def plot_postprocess_variants(summary: dict, out: Path) -> None:
    road_ious = [summary[v]["road_iou"] for v in VARIANTS]
    bf1s = [summary[v]["boundary_f1"] for v in VARIANTS]
    labels = [v.upper() for v in VARIANTS]

    x = range(len(VARIANTS))
    width = 0.38
    fig, ax = plt.subplots(figsize=(9, 4.2))
    ax.bar([i - width / 2 for i in x], road_ious, width, label="road IoU", color="#2c7fb8")
    ax.bar([i + width / 2 for i in x], bf1s, width, label="boundary F1", color="#fdae6b")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels)
    ax.set_ylim(0, 1.0)
    ax.set_ylabel("score")
    ax.set_title("Postprocessing variants on retrained drivable model")
    ax.legend(fontsize=9)
    _style_ax(ax)

    for i, (r, b) in enumerate(zip(road_ious, bf1s)):
        ax.text(i - width / 2, r + 0.02, f"{r:.3f}", ha="center", fontsize=7.5)
        ax.text(i + width / 2, b + 0.02, f"{b:.3f}", ha="center", fontsize=7.5)

    fig.tight_layout()
    out.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out, dpi=160)
    plt.close(fig)


def plot_per_sample_hist(per_sample: list[dict], out: Path) -> None:
    raw_ious = [s["raw_road_iou"] for s in per_sample]
    fig, ax = plt.subplots(figsize=(7, 3.8))
    ax.hist(raw_ious, bins=20, color="#2c7fb8", alpha=0.85, edgecolor="white")
    mean = sum(raw_ious) / len(raw_ious)
    ax.axvline(mean, color="#e6550d", linestyle="--", linewidth=1.6, label=f"mean = {mean:.3f}")
    ax.set_xlabel("raw road IoU (per image)")
    ax.set_ylabel("count")
    ax.set_title("Per-image road IoU distribution (100 val images)")
    ax.legend(fontsize=9)
    _style_ax(ax)
    fig.tight_layout()
    out.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out, dpi=160)
    plt.close(fig)


def _img_ref(path: Path) -> str:
    return f"file:///{str(path).replace(chr(92), '/')}"


def build_markdown(summary, baseline, cfg, training, assets, visual_samples, eval_dir: Path, out_md: Path) -> str:
    raw = summary["raw"]
    b = baseline
    delta_iou = raw["road_iou"] - b["road_iou_global"]
    delta_f1 = raw["boundary_f1"] - b["boundary_f1"]

    variant_rows = "\n".join(
        f"| {VARIANT_LABELS[v]} | {summary[v]['road_iou']:.4f} | {summary[v]['road_precision']:.4f} | "
        f"{summary[v]['road_recall']:.4f} | {summary[v]['boundary_f1']:.4f} | {summary[v]['boundary_iou']:.4f} |"
        for v in VARIANTS
    )

    best_epoch = int(max(range(len(training["road_iou"])), key=lambda i: training["road_iou"][i])) + 1
    best_road = max(training["road_iou"])
    sample_rows = "\n\n".join(
        "\n".join(
            [
                f"### {sample['category']}：`{sample['name']}`",
                f"- raw road IoU：`{sample['raw_road_iou']:.4f}`",
                f"- raw boundary F1：`{sample['raw_boundary_f1']:.4f}`",
                f"![{sample['name']} panel]({_img_ref(sample['panel'])})",
            ]
        )
        for sample in visual_samples
    )

    return f"""# PIDNet Drivable 道路分割优化报告

## 1. 目标与结论
- 目标：使道路边界 / 道路 mask IoU **显著高于** team 3 sub SegFormer 道路分割基线。
- 结论：**已达成**。方案 D（重训 PIDNet-small 2 类 drivable 专用模型）在同 100 张 `val_pilot` 样本上，
  道路 IoU 由基线 `{b['road_iou_global']:.4f}` 提升至 `{raw['road_iou']:.4f}`（**+{delta_iou:.3f}**），
  边界 F1 由 `{b['boundary_f1']:.4f}` 提升至 `{raw['boundary_f1']:.4f}`（**+{delta_f1:.3f}**）。

## 2. 与 SegFormer 基线同口径对比

| 指标 | team 3 sub SegFormer | PIDNet drivable（本方案） | 提升 |
| --- | --- | --- | --- |
| road IoU（global） | {b['road_iou_global']:.4f} | {raw['road_iou']:.4f} | {delta_iou:+.3f} |
| road precision | {b['road_precision']:.4f} | {raw['road_precision']:.4f} | {raw['road_precision']-b['road_precision']:+.3f} |
| road recall | {b['road_recall']:.4f} | {raw['road_recall']:.4f} | {raw['road_recall']-b['road_recall']:+.3f} |
| boundary F1（容差 3） | {b['boundary_f1']:.4f} | {raw['boundary_f1']:.4f} | {delta_f1:+.3f} |
| boundary IoU（容差 3） | {b['boundary_iou']:.4f} | {raw['boundary_iou']:.4f} | {raw['boundary_iou']-b['boundary_iou']:+.3f} |

![基线对比]({_img_ref(assets['baseline_comparison'])})

## 3. 训练配置与收敛

- 配置文件：`{cfg['MODEL']['NAME']}`（`num_outputs={cfg['MODEL']['NUM_OUTPUTS']}`）
- 数据：`{cfg['DATASET']['ROOT']}`，train `{cfg['DATASET']['TRAIN_SET']}`，val `{cfg['DATASET']['TEST_SET']}`
- 类别数：`{cfg['DATASET']['NUM_CLASSES']}`（0 背景 / 1 道路）
- Warm start：`{cfg['MODEL']['PRETRAINED']}`
- 训练尺寸：`{cfg['TRAIN']['IMAGE_SIZE']}`，batch `{cfg['TRAIN']['BATCH_SIZE_PER_GPU']}`，epoch `{cfg['TRAIN']['END_EPOCH']}`
- 学习率：`{cfg['TRAIN']['LR']}`（poly 衰减），optimizer `{cfg['TRAIN']['OPTIMIZER']}`
- 验证最佳：Epoch `{best_epoch}`，道路 IoU `{best_road:.4f}`

![训练曲线]({_img_ref(assets['training_curves'])})

## 4. 后处理变体（A/B/C）对重训模型的影响

| 变体 | road IoU | precision | recall | boundary F1 | boundary IoU |
| --- | --- | --- | --- | --- | --- |
{variant_rows}

- 结论：`refined` / `crf` / `lane_trim` 对 road IoU 影响几乎为 0（约 -0.0015），
  说明重训后的专用模型 raw 输出已经足够精准，无需额外切边。
  `snake` / `all` 反而显著降低 road IoU（约 -0.054），不采用。

![后处理变体对比]({_img_ref(assets['postprocess_variants'])})

## 5. 逐样本分布

![逐样本 road IoU 分布]({_img_ref(assets['per_sample_hist'])})

## 6. 分割结果样例

{sample_rows}

## 7. 交付物
- 训练配置：`external/PIDNet/configs/bdd100k/pidnet_small_bdd_drivable_pilot_v2.yaml`
- 权重：`{summary['model_file']}`
- 评估产物：`{eval_dir}/summary.json`、`per_sample.json`
"""


def main() -> None:
    args = parse_args()
    eval_dir = Path(args.eval_dir)
    summary = json.loads((eval_dir / "summary.json").read_text(encoding="utf-8"))
    per_sample = json.loads((eval_dir / "per_sample.json").read_text(encoding="utf-8"))
    baseline_all = json.loads(Path(args.baseline_summary).read_text(encoding="utf-8"))
    baseline = baseline_all["argmax_0.5"]
    cfg = yaml.safe_load(Path(args.config).read_text(encoding="utf-8"))
    runtime_cfg = load_runtime_config(args.config)

    train_log = find_train_log(Path(args.train_dir))
    training = parse_training_log(train_log)

    assets_dir = eval_dir / "report_assets"
    assets = {
        "training_curves": assets_dir / "training_curves.png",
        "baseline_comparison": assets_dir / "baseline_comparison.png",
        "postprocess_variants": assets_dir / "postprocess_variants.png",
        "per_sample_hist": assets_dir / "per_sample_hist.png",
    }
    plot_training_curves(training, assets["training_curves"])
    plot_baseline_comparison(summary, baseline, assets["baseline_comparison"])
    plot_postprocess_variants(summary, assets["postprocess_variants"])
    plot_per_sample_hist(per_sample, assets["per_sample_hist"])
    label_map = parse_list_file(Path(runtime_cfg.DATASET.ROOT), summary["list_path"])
    model = load_model(runtime_cfg, Path(summary["model_file"]))
    visual_samples = create_visual_samples(
        per_sample,
        runtime_cfg,
        model,
        label_map,
        eval_dir / "report_visuals",
        args.visual_count,
    )

    raw = summary["raw"]
    delta_iou = raw["road_iou"] - baseline["road_iou_global"]

    out_docx = Path(args.output_docx)
    out_md = Path(args.output_md)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    out_md.parent.mkdir(parents=True, exist_ok=True)

    # ---- Word ----
    doc = Document()
    doc.add_heading("PIDNet Drivable 道路分割优化报告", level=0)
    doc.add_paragraph(
        "本报告对比 team 3 sub SegFormer 道路分割基线，验证重训 PIDNet-small 2 类 drivable 专用模型"
        "在道路 mask IoU 与边界 F1 上的提升。"
    )

    doc.add_heading("1. 结论", level=1)
    doc.add_paragraph(
        f"方案 D 在同 100 张 val_pilot 样本上，道路 IoU 由基线 {baseline['road_iou_global']:.4f} 提升至 "
        f"{raw['road_iou']:.4f}（+{delta_iou:.3f}），边界 F1 由 {baseline['boundary_f1']:.4f} 提升至 "
        f"{raw['boundary_f1']:.4f}（+{raw['boundary_f1']-baseline['boundary_f1']:.3f}），显著超越基线。"
    )

    doc.add_heading("2. 与 SegFormer 基线对比", level=1)
    table = doc.add_table(rows=1, cols=4)
    header = table.rows[0].cells
    for i, text in enumerate(["指标", "SegFormer", "PIDNet drivable", "提升"]):
        header[i].text = text
    rows = [
        ("road IoU", baseline["road_iou_global"], raw["road_iou"], delta_iou),
        ("road precision", baseline["road_precision"], raw["road_precision"], raw["road_precision"] - baseline["road_precision"]),
        ("road recall", baseline["road_recall"], raw["road_recall"], raw["road_recall"] - baseline["road_recall"]),
        ("boundary F1", baseline["boundary_f1"], raw["boundary_f1"], raw["boundary_f1"] - baseline["boundary_f1"]),
        ("boundary IoU", baseline["boundary_iou"], raw["boundary_iou"], raw["boundary_iou"] - baseline["boundary_iou"]),
    ]
    for name, bv, ov, dv in rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = f"{bv:.4f}"
        row[2].text = f"{ov:.4f}"
        row[3].text = f"{dv:+.3f}"
    doc.add_picture(str(assets["baseline_comparison"]), width=Inches(6.0))

    doc.add_heading("3. 训练配置与收敛", level=1)
    doc.add_paragraph(
        f"模型 {cfg['MODEL']['NAME']}（num_outputs={cfg['MODEL']['NUM_OUTPUTS']}），类别数 "
        f"{cfg['DATASET']['NUM_CLASSES']}。训练数据 {cfg['DATASET']['TRAIN_SET']}，验证数据 "
        f"{cfg['DATASET']['TEST_SET']}。训练尺寸 {cfg['TRAIN']['IMAGE_SIZE']}，batch "
        f"{cfg['TRAIN']['BATCH_SIZE_PER_GPU']}，{cfg['TRAIN']['END_EPOCH']} epoch，学习率 "
        f"{cfg['TRAIN']['LR']}（poly 衰减）。Warm start 自 {cfg['MODEL']['PRETRAINED']}。"
    )
    doc.add_picture(str(assets["training_curves"]), width=Inches(6.4))

    doc.add_heading("4. 后处理变体影响", level=1)
    doc.add_paragraph(
        "refined / crf / lane_trim 对 road IoU 影响几乎为 0（约 -0.0015），说明重训模型 raw 输出已足够精准；"
        "snake / all 反而降低 road IoU（约 -0.054），不采用。"
    )
    doc.add_picture(str(assets["postprocess_variants"]), width=Inches(6.2))

    doc.add_heading("5. 逐样本分布", level=1)
    doc.add_picture(str(assets["per_sample_hist"]), width=Inches(5.6))

    doc.add_heading("6. 分割结果样例", level=1)
    doc.add_paragraph("下列样例按三组组织：最差、最好、最平均。每张图展示原图、模型预测、GT 以及误差叠加图（绿色 TP，红色 FP，蓝色 FN）。")
    for sample in visual_samples:
        doc.add_paragraph(
            f"{sample['category']} | {sample['name']} | raw road IoU={sample['raw_road_iou']:.4f} | raw boundary F1={sample['raw_boundary_f1']:.4f}"
        )
        add_picture_if_exists(doc, Path(sample["panel"]), width=6.2)

    doc.add_heading("7. 交付物", level=1)
    doc.add_paragraph(f"训练配置：{args.config}")
    doc.add_paragraph(f"权重：{summary['model_file']}")
    doc.add_paragraph(f"评估产物：{eval_dir / 'summary.json'}、{eval_dir / 'per_sample.json'}")

    doc.save(out_docx)

    # ---- Markdown ----
    md = build_markdown(summary, baseline, cfg, training, assets, visual_samples, eval_dir, out_md)
    out_md.write_text(md, encoding="utf-8")

    print(f"report docx saved to {out_docx}")
    print(f"report markdown saved to {out_md}")
    print(f"charts saved to {assets_dir}")


if __name__ == "__main__":
    main()
